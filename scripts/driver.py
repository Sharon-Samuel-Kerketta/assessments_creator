import os
from selenium import webdriver
from webdriver_manager.chrome import ChromeDriverManager
from selenium.webdriver.common.by import By
import docx

from time import sleep

from scripts.utils import filter_text, sanitize_filename


def setup_browser(form_link:str) -> webdriver.Chrome:
    browser = webdriver.Chrome(ChromeDriverManager().install())
    browser.get(form_link)
    sleep(5)

    return browser


def get_raw_qno(browser:webdriver.Chrome) -> list:
    question_answer_list = []

    raw_qno = browser.find_elements(By.CLASS_NAME, "geS5n")
    raw_qno = [qno.text for qno in raw_qno]
    return raw_qno


def filter_qno(raw_qno:list) -> tuple[list, list]:
    questions = []
    options = []
    for text in raw_qno:
        q, a = filter_text(text)
        questions.append(q)
        options.append(a)

    return questions, options


def create_docx(questions:list, options:list, output_file:str = "assessment.docx") -> None:

    doc = docx.Document()

    for question, options in zip(questions,options):
        doc.add_paragraph(question)
        for option in options:
            doc.add_paragraph(option)
        doc.add_paragraph("\n")

    os.makedirs("data", exist_ok=True)
    doc.save(os.path.join(os.path.curdir, "data", output_file))


def scrape(form_link:str, output_file:str) -> None:

    print("\n\nSit back and relax while your form is getting scraped ......")

    # setup browser
    browser = setup_browser(form_link)

    # Find the div with questions and options
    raw_qno = get_raw_qno(browser)

    if raw_qno:
        # filter to get only the questions and options 
        questions, options = filter_qno(raw_qno)

        # create a docx file with the extracted question and options
        create_docx(questions, options, output_file)

        print(f"Scraping done successfully! Check {output_file}\nThank you! \\m/")

    else:
        print("No data found! Please check the link provided")